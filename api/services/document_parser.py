"""Document parsing service — PDF, DOCX, TXT, MD to raw text.

Handles German characters (äöüß) correctly across all formats.
All functions are async; blocking I/O runs in asyncio.to_thread.
"""

import asyncio
from pathlib import Path

import structlog

logger = structlog.get_logger()

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


async def parse_document(file_path: str, filename: str) -> str:
    """Extract raw text from a document file.

    Supports .pdf (PyMuPDF), .docx (python-docx), .txt and .md (UTF-8).
    German characters (äöüß) are preserved in all formats.

    Args:
        file_path: Absolute path to the uploaded file on disk.
        filename: Original filename (used to detect extension).

    Returns:
        Extracted raw text as a single string.

    Raises:
        ValueError: If the file type is not supported.
        IOError: If the file cannot be read.
    """
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}'. Accepted: {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    logger.info("document_parse_start", filename=filename, ext=ext)

    if ext == ".pdf":
        text = await asyncio.to_thread(_parse_pdf, file_path)
    elif ext == ".docx":
        text = await asyncio.to_thread(_parse_docx, file_path)
    else:  # .txt or .md
        text = await asyncio.to_thread(_parse_text, file_path)

    logger.info("document_parse_done", filename=filename, chars=len(text))
    return text


def _parse_pdf(file_path: str) -> str:
    """Extract text from PDF using PyMuPDF (fitz) + pdfplumber for tables.

    Strategy:
      1. fitz extracts running prose text per page.
      2. pdfplumber extracts tables (compliance evidence often lives in tables).
         Tables are appended after the prose as tab-separated rows.
      3. If both extractors produce < 100 chars (scanned/image PDF), falls back
         to pytesseract OCR on rasterised page images.

    German ligatures and special characters handled natively by fitz.
    """
    import fitz  # PyMuPDF

    doc = fitz.open(file_path)
    page_parts: list[str] = []

    for page_idx, page in enumerate(doc):
        prose = page.get_text("text").strip()
        page_parts.append(prose)

    doc.close()

    # ── Table extraction via pdfplumber ───────────────────────────────────────
    table_parts: list[str] = []
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                for table in page.extract_tables():
                    if not table:
                        continue
                    rows = []
                    for row in table:
                        cells = [str(c).strip() if c else "" for c in row]
                        row_text = "\t".join(c for c in cells if c)
                        if row_text:
                            rows.append(row_text)
                    if rows:
                        table_parts.append("\n".join(rows))
    except ImportError:
        pass  # pdfplumber optional — prose-only extraction is fine
    except Exception:
        pass  # malformed tables are not fatal

    full_text = "\n\n".join(page_parts)
    if table_parts:
        full_text = full_text + "\n\n" + "\n\n".join(table_parts)

    # ── OCR fallback for scanned PDFs ─────────────────────────────────────────
    if len(full_text.strip()) < 100:
        full_text = _ocr_pdf(file_path) or full_text

    return full_text


def _ocr_pdf(file_path: str) -> str | None:
    """Rasterise PDF pages and run pytesseract OCR. Returns None if unavailable."""
    try:
        import fitz
        import pytesseract
        from PIL import Image
        import io
    except ImportError:
        return None  # pytesseract + Pillow are optional

    try:
        doc = fitz.open(file_path)
        pages_text: list[str] = []
        for page in doc:
            # Render at 200 DPI for reliable OCR
            pix = page.get_pixmap(dpi=200)
            img = Image.open(io.BytesIO(pix.tobytes("png")))
            text = pytesseract.image_to_string(img, lang="eng+deu")
            pages_text.append(text.strip())
        doc.close()
        return "\n\n".join(pages_text)
    except Exception:
        return None


def _parse_docx(file_path: str) -> str:
    """Extract text from DOCX using python-docx.

    Iterates all paragraphs; preserves paragraph breaks.
    Tables are extracted row by row with tab separation.
    """
    from docx import Document

    doc = Document(file_path)
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Also extract table content
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n\n".join(parts)


def _parse_text(file_path: str) -> str:
    """Read plain text or Markdown file as UTF-8."""
    with open(file_path, encoding="utf-8", errors="replace") as f:
        return f.read()
